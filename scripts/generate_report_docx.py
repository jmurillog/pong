"""Generate REPORT.docx from REPORT.md content with simple academic formatting."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIGURES = os.path.join(BASE, "outputs", "ppo_pong", "analysis")
OUT = os.path.join(BASE, "REPORT.docx")

doc = Document()

# ── Page margins (2.5 cm all round) ──────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    sizes = {1: 13, 2: 12, 3: 11}
    set_font(run, size=sizes.get(level, 11), bold=True)
    return p

def add_para(text, italic=False, space_before=0, space_after=6, first_line=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(18)
    # Handle **bold** inline
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        set_font(run, bold=(i % 2 == 1), italic=italic)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3 + 0.2 * level)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        set_font(run, bold=(i % 2 == 1))
    return p

def add_figure(filename, caption):
    path = os.path.join(FIGURES, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run()
        run.add_picture(path, width=Inches(5.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    set_font(r, size=10, italic=True)

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        set_font(r, size=10, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Light grey shading
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(str(val))
            set_font(r, size=10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacing after table

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)

def separator():
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ═══════════════════════════════════════════════════════════════════════════════
#  TITLE & AUTHORS
# ═══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(6)
r = title_p.add_run(
    "Can a PPO Reinforcement Learning Agent Reach and Exceed a\n"
    "Novice Human Baseline in Pong Under a Fixed Compute Budget?"
)
set_font(r, size=15, bold=True)

authors_p = doc.add_paragraph()
authors_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors_p.paragraph_format.space_after = Pt(2)
r = authors_p.add_run(
    "José María Murillo, Hans Helmrich, Beatriz Wahle, Valentín Miguel, Pablo Chen, Nicolás Cubillo"
)
set_font(r, size=11)

inst_p = doc.add_paragraph()
inst_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
inst_p.paragraph_format.space_after = Pt(16)
r = inst_p.add_run("IE University · Emerging Topics in Data Analysis · April 2026")
set_font(r, size=10, italic=True)

# ═══════════════════════════════════════════════════════════════════════════════
#  ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Abstract", level=1)
add_para(
    "We establish a reproducible benchmark comparing Proximal Policy Optimization (PPO) against a "
    "measured novice human baseline on Atari Pong. Across five random seeds, PPO trained for "
    "1,000,000 timesteps achieves a mean episodic return of −8.46 (95% CI: [−10.78, −6.58]), already "
    "exceeding the human baseline mean of −15.73 (95% CI: [−16.40, −15.06]) by 7.27 points "
    "(CI: [4.98, 9.27]). Extending training to 5,000,000 timesteps raises agent performance to +15.61 "
    "(CI: [+14.33, +16.62]), a 31.34-point margin over humans, with 100% win rate across all seeds. "
    "Analysis of learning curves reveals PPO first exceeds the human mean at approximately 800,000 "
    "timesteps. Seed variance decreases from σ = 2.64 at 1M to σ = 1.47 at 5M steps, confirming "
    "reliable convergence. All code and raw data are publicly available."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Introduction", level=1)
add_para(
    "Reinforcement learning (RL) has produced landmark results in artificial intelligence — from "
    "AlphaGo defeating world champions (Silver et al., 2016) to DQN mastering Atari games from raw "
    "pixels (Mnih et al., 2015). Yet most of this work operates at scales that are inaccessible to "
    "teams without large compute budgets. A practically important but underexplored question follows: "
    "how far can a standard RL algorithm get on a well-defined task under a fixed, modest compute "
    "budget — and at what point does it cross the threshold of human competence?"
)

add_para(
    "**Research gap.** Most RL benchmarking studies compare algorithms against each other or against "
    "human scores that were collected under non-standardized conditions (Mnih et al., 2015 used "
    "professional game testers with unlimited practice). Critically, these comparisons rarely quantify "
    "uncertainty: they report point estimates without confidence intervals, making it impossible to "
    "assess whether a reported performance gap is reliable or a product of random seed selection. "
    "Henderson et al. (2018) demonstrated that different seeds, codebases, and hyperparameter choices "
    "can produce dramatically different results for the same algorithm on the same task — a finding "
    "that directly motivates our multi-seed design. Agarwal et al. (2021) further showed that the "
    "field has systematically underreported uncertainty and advocated for bootstrap confidence intervals "
    "as the appropriate evaluation protocol for deep RL."
)

add_para(
    "**Novelty.** We address this gap by: (1) collecting a fresh, in-house novice human baseline under "
    "the same evaluation conditions as the agent (same environment, opponent, and scoring system); "
    "(2) reporting bootstrap 95% confidence intervals for the agent mean, the human mean, and the "
    "agent−human difference simultaneously, enabling statistically principled comparison; and "
    "(3) framing the study around two concrete hypotheses with a pre-specified success criterion. "
    "The primary stakeholder is any small engineering team that must decide whether PPO is a sensible "
    "first choice for game agent development under time and compute constraints: our sample-efficiency "
    "curve and seed-stability estimates directly support that decision."
)

add_para(
    "Beyond this specific experiment, the approach itself — measuring a human baseline under "
    "controlled conditions and reporting uncertainty at every level of the comparison — offers a "
    "replicable evaluation methodology that addresses long-standing validity concerns in RL "
    "benchmarking (Agarwal et al., 2021). The field has protocols for comparing algorithms to each "
    "other; it lacks a standard for comparing them to humans with appropriate uncertainty "
    "quantification. This study contributes one such protocol."
)

add_para(
    "**Hypotheses.** H1: PPO will reach or exceed the novice human baseline within 5M environment "
    "steps. H2: Performance will show meaningful variance across seeds, such that single-run reporting "
    "would meaningfully misrepresent algorithm reliability (Henderson et al., 2018). Both hypotheses "
    "were tested prospectively and their confirmation assessed against bootstrap CIs."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  RESULTS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Results", level=1)

add_heading("Final performance and human comparison", level=2)
add_para(
    "Table 1 summarizes final agent performance for both budgets alongside the measured human "
    "baseline, with bootstrap 95% confidence intervals."
)

# Table 1
cap = doc.add_paragraph()
cap.paragraph_format.space_before = Pt(8)
cap.paragraph_format.space_after = Pt(4)
r = cap.add_run("Table 1. Final performance: PPO agent vs. novice human baseline (bootstrap 95% CIs, 10,000 resamples)")
set_font(r, size=10, bold=True)

add_table(
    headers=["Condition", "n", "Mean Return", "SE", "95% CI", "Win Rate"],
    rows=[
        ["Human baseline",    "90 ep.",  "−15.73", "0.35", "[−16.40, −15.06]", "0%"],
        ["PPO — 1M steps",    "5 seeds", "−8.46",  "1.18", "[−10.78, −6.58]",  "10%"],
        ["PPO — 5M steps",    "5 seeds", "+15.61", "0.66", "[+14.33, +16.62]", "100%"],
    ]
)

add_para(
    "**H1 confirmed.** PPO exceeds the novice human baseline at both budget levels. The 95% bootstrap "
    "CI for the agent−human difference at 1M steps is [+4.98, +9.27], entirely excluding zero. At 5M "
    "steps, the difference grows to +31.34 (CI: [+29.94, +32.57]), with all five seeds achieving 100% "
    "win rate."
)
add_para(
    "**H2 confirmed.** At 1M steps, per-seed final returns span [−12.35, −6.20] with σ = 2.64: a "
    "cherry-picked best run appears 50% better than the worst, confirming that single-seed reporting "
    "is meaningfully misleading. At 5M steps, variance shrinks (σ = 1.47, range [+13.25, +17.00]), "
    "but seed-level differences remain non-trivial for compute-constrained decisions."
)
add_para(
    "Figure 2 shows per-seed final returns for both budgets. All five 5M seeds individually exceed "
    "the entire human distribution. All five 1M seeds individually exceed the human mean."
)

add_figure("seed_variance.png", "Figure 2. Per-seed final returns for 1M and 5M budgets with human baseline reference (dashed).")

add_heading("Learning dynamics", level=2)
add_para(
    "Figure 1 shows learning curves (mean ± 1 SE across five seeds) for both budgets, alongside the "
    "human baseline, revealing how the agent arrives at the headline numbers above. Both budget "
    "trajectories share a cold-start phase: mean return is −21.0 for the first 200,000–400,000 "
    "timesteps, indicating the agent has not yet discovered it can score any points. Around step "
    "500,000, performance rises sharply. At 800,000 timesteps — approximately 80% into the shorter "
    "budget — the cross-seed mean first exceeds the human baseline mean of −15.73, directly answering "
    "RQ1: PPO surpasses novice human performance at roughly 800k training steps, well within a 1M-step "
    "budget."
)
add_para(
    "The 5M trajectory continues improving after the 1M run terminates. The agent crosses zero return "
    "at approximately 1.4M steps, reaches a 100% win rate by 3.6M steps, and stabilizes at +15.61 "
    "return at 5M steps. This characteristic two-phase structure — rapid acquisition of game mechanics "
    "(500k–1.5M) followed by gradual refinement — is consistent across all five seeds."
)

add_figure("learning_curves.png", "Figure 1. Learning curves (mean ± 1 SE across five seeds) for 1M and 5M budgets. Human baseline mean shown as dashed line.")

# ═══════════════════════════════════════════════════════════════════════════════
#  DISCUSSION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Discussion", level=1)

add_heading("Interpretation", level=2)
add_para(
    "The most striking finding is the speed of human-level surpassing: approximately 800,000 timesteps, "
    "corresponding to under two hours of wall-clock training on a consumer laptop with GPU acceleration. "
    "This is far earlier than the 1M-step budget would suggest is needed, meaning teams operating under "
    "tight compute constraints can still expect above-human performance. The cold-start plateau "
    "(steps 1–400k) followed by rapid improvement is consistent with the well-documented reward sparsity "
    "problem in Atari: the agent must first discover that scoring is possible before systematic "
    "improvement can occur."
)
add_para(
    "At 5M steps, the 100% win rate across all seeds indicates the agent has effectively solved Pong at "
    "the difficulty of the built-in ROM opponent. Remaining episodic variation (+13 to +17) reflects "
    "game stochasticity, not policy uncertainty."
)

add_heading("Implications and recommendations", level=2)
add_para(
    "For practitioners, the results support the following decision framework: a 1M-step PPO run is a "
    "low-risk choice for reaching above-human novice performance on Pong, while a 5M-step run achieves "
    "near-perfect play with high reliability. The seed-stability improvement from 1M to 5M (σ: 2.64 → "
    "1.47) suggests that increasing compute not only raises the mean but also reduces the risk of a poor "
    "outcome from unlucky initialization. We recommend always running at minimum 3–5 seeds before "
    "reporting results, as single-run performance estimates carry substantial uncertainty even in a "
    "relatively easy environment like Pong."
)

add_heading("Ethical considerations", level=2)
add_para(
    "While this study is technically benign, it touches on broader considerations relevant to RL "
    "benchmarking. Human-AI comparisons carry implicit claims about substitutability: framing an agent "
    "as 'superhuman' after comparing it to novice players under non-representative conditions would be "
    "misleading. Our study deliberately limits its claims — we benchmark against novice humans playing "
    "under time pressure with unfamiliar controls, not experienced players or game-specific experts. "
    "Additionally, the fabricated nature of the Atari environment means the agent's superiority reflects "
    "mastery of a fixed, known-dynamics simulator, not general intelligence. These distinctions matter "
    "when communicating RL results to non-technical audiences or policymakers."
)

add_heading("Limitations", level=2)
add_para(
    "The most significant limitation is a methodological validity threat in the human-agent comparison "
    "itself. The PPO agent observes 84×84 grayscale pixel frames directly from the ALE ROM environment, "
    "while human participants played via a web-based browser game built to mirror the same opponent and "
    "scoring rules. These are not identical perceptual environments: the human sees a rendered, colour "
    "interface at variable frame rates and uses keyboard inputs, while the agent processes a canonical "
    "pixel stack with precise temporal alignment. This confound means the comparison is not fully "
    "controlled — superior agent performance could partly reflect differences in interface familiarity "
    "or sensorimotor latency, not only policy quality. Partially mitigating this concern is that human "
    "Pong performance is predominantly limited by opponent difficulty and game understanding, not "
    "reaction-time precision at typical game speeds; nevertheless, this confound should be considered "
    "when generalising the results."
)
add_para(
    "Second, the human participants are novice players (no prior Pong experience); a baseline including "
    "experienced players would set a higher bar and likely delay the crossover point. Third, Pong is one "
    "of the easier Atari games for deep RL — the same budget on games such as Montezuma's Revenge or "
    "Pitfall would yield dramatically different results, and the crossover timing should not be assumed "
    "to generalise across games. Fourth, the agent is trained against a fixed, deterministic ROM "
    "opponent; its policy may not transfer to adaptive or human opponents. Fifth, wall-clock training "
    "time was not systematically recorded across all seeds, limiting reproducibility for "
    "time-constrained practitioners."
)

add_heading("Future directions", level=2)
add_para("Natural extensions include:")
add_bullet("testing the same pipeline on harder Atari games to characterize where the human-surpassing threshold lies as a function of game complexity;")
add_bullet("collecting an experienced human baseline to establish an upper bound that RL may not yet reach;")
add_bullet("comparing PPO against DQN and A3C under identical budgets and the same human baseline to isolate algorithmic differences from evaluation differences;")
add_bullet("investigating whether the 800k crossover point is robust to hyperparameter changes, particularly the clip range and learning rate.")

# ═══════════════════════════════════════════════════════════════════════════════
#  METHODS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Methods", level=1)

add_heading("Environment and preprocessing", level=2)
add_para(
    "All experiments use ALE/Pong-v5 via Gymnasium 1.0.0 and ale-py 0.10.1. Standard Atari "
    "preprocessing is applied: frames are converted to grayscale and resized to 84×84 pixels; four "
    "consecutive frames are stacked to provide temporal information to the policy; rewards are clipped "
    "to [−1, +1] during training to normalize the gradient signal. Eight parallel environments are "
    "used throughout training."
)

add_heading("Algorithm", level=2)
add_para(
    "We use Proximal Policy Optimization (Schulman et al., 2017) implemented in Stable-Baselines3 "
    "2.4.1. PPO constrains each policy gradient update to remain within a trust region via a clipped "
    "surrogate objective, balancing sample efficiency with training stability. Hyperparameters follow "
    "the Stable-Baselines3 Atari defaults: n_steps = 128, n_epochs = 4, mini-batch size = 256, clip "
    "range ε = 0.1, learning rate α = 2.5 × 10⁻⁴, discount factor γ = 0.99, GAE parameter λ = 0.95, "
    "entropy coefficient = 0.01, value function coefficient = 0.5, max gradient norm = 0.5. No "
    "hyperparameter tuning was performed; the goal is to assess what a default PPO run achieves, which "
    "is directly consistent with our stakeholder question."
)

add_heading("Experimental design", level=2)
add_para(
    "Two training budgets were evaluated: 1,000,000 and 5,000,000 timesteps. Each was run with five "
    "independent random seeds {11, 22, 33, 44, 55}, yielding 10 total training runs. Every 100,000 "
    "timesteps, the policy was evaluated deterministically (no exploration) on 20 episodes using a "
    "fixed set of evaluation seeds {1000, …, 1019}, ensuring checkpoint comparisons are not confounded "
    "by evaluation stochasticity. The primary metric is mean episodic return; win rate (fraction of "
    "episodes with positive return) is the secondary metric."
)

add_heading("Human baseline", level=2)
add_para(
    "Six participants (IE University students, no prior Pong experience) each played 15 complete "
    "episodes of Pong against the built-in ROM opponent, yielding 90 total episodes. Episode return is "
    "defined as (player points) − (opponent points), range [−21, +21]. Data were recorded using the "
    "project web dashboard."
)

add_heading("Statistical analysis", level=2)
add_para(
    "Agent performance is summarized as the mean of seed-level final returns (n = 5) ± standard error "
    "(SE = σ/√n). Human performance is the mean of all episode returns (n = 90) ± SE. Bootstrap 95% "
    "CIs are computed for: (a) agent mean, (b) human mean, and (c) the agent−human difference. All "
    "bootstrap estimates use 10,000 resamples with a fixed RNG seed (12345) for reproducibility. The "
    "success criterion was pre-specified as: the 5M agent mean exceeds the human mean with a bootstrap "
    "CI for the difference that entirely excludes zero."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  AUTHOR CONTRIBUTIONS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Author Contributions", level=1)

cap2 = doc.add_paragraph()
cap2.paragraph_format.space_after = Pt(4)
r = cap2.add_run("Table 2. Author contributions")
set_font(r, size=10, bold=True)

add_table(
    headers=["Name", "Role", "Contribution"],
    rows=[
        ["José María Murillo", "Project Lead",            "Overall coordination, PPO training pipeline, hyperparameter setup, final report editing"],
        ["Hans Helmrich",      "RL Engineer",             "Environment setup (Gymnasium/Atari), frame preprocessing, reward logging, web dashboard"],
        ["Beatriz Wahle",      "Data & Evaluation",       "Human baseline collection protocol, evaluation metrics, seed variance analysis"],
        ["Valentín Miguel",    "Literature & Writing",    "Literature review, methods section writing, references management"],
        ["Pablo Chen",         "Experiments & Analysis",  "Multi-seed experiment runs, learning curve visualization, statistical reporting"],
        ["Nicolás Cubillo",    "Reproducibility & Code",  "Code documentation, reproducibility pipeline, appendix, AI usage statement"],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  APPENDIX — REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Appendix", level=1)
add_heading("References", level=2)

refs = [
    "Agarwal, R., Schwarzer, M., Castro, P. S., Courville, A. C., & Bellemare, M. (2021). Deep reinforcement learning at the edge of the statistical precipice. Advances in Neural Information Processing Systems, 34, 29304–29320.",
    "Bellemare, M. G., Naddaf, Y., Veness, J., & Bowling, M. (2013). The Arcade Learning Environment: An evaluation platform for general agents. Journal of Artificial Intelligence Research, 47, 253–279.",
    "Henderson, P., Islam, R., Bachman, P., Pineau, J., Precup, D., & Meger, D. (2018). Deep reinforcement learning that matters. Proceedings of the AAAI Conference on Artificial Intelligence, 32(1).",
    "Mnih, V., Kavukcuoglu, K., Silver, D., et al. (2015). Human-level control through deep reinforcement learning. Nature, 518(7540), 529–533.",
    "Raffin, A., Hill, A., Gleave, A., Kanervisto, A., Ernestus, M., & Dormann, N. (2021). Stable-Baselines3: Reliable reinforcement learning implementations. Journal of Machine Learning Research, 22(268), 1–8.",
    "Schulman, J., Wolski, F., Dhariwal, P., Radford, A., & Klimov, O. (2017). Proximal policy optimization algorithms. arXiv preprint arXiv:1707.06347.",
    "Silver, D., Huang, A., Maddison, C. J., et al. (2016). Mastering the game of Go with deep neural networks and tree search. Nature, 529(7587), 484–489.",
]
for ref in refs:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(ref)
    set_font(run, size=10, italic=False)

# ═══════════════════════════════════════════════════════════════════════════════
#  APPENDIX — CODE AVAILABILITY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Code Availability Statement", level=2)
add_para(
    "All code is available in the project repository. The pipeline is structured as follows."
)

add_para("**Environment setup:**")
add_code("python3 -m venv .venv && source .venv/bin/activate")
add_code("pip install -r requirements.txt")
add_code("AutoROM --accept-license")

add_para("**Reproduce the full 10-run experiment matrix:**")
add_code(
    "python scripts/run_experiments.py \\\n"
    "  --env-id ALE/Pong-v5 --budgets 1000000,5000000 \\\n"
    "  --seeds 11,22,33,44,55 --eval-freq 100000 \\\n"
    "  --n-eval-episodes 20 --n-envs 8 --output-root outputs --skip-if-complete"
)

add_para("**Regenerate all tables, figures, and statistics:**")
add_code(
    "python scripts/analyze_results.py --output-root outputs \\\n"
    "  --experiment-name ppo_pong --human-csv data/human_baseline/human_returns.csv"
)

add_para(
    "Raw per-run CSV logs, model checkpoints, and runtime metadata are stored in "
    "outputs/ppo_pong/{budget}/seed_{seed}/. Training at full scale on Apple M3 with MPS acceleration "
    "takes approximately 8–10 hours per 5M-step seed."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  APPENDIX — AI USAGE STATEMENT
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("AI Usage Statement", level=2)
add_para(
    "This project used AI tooling (Claude Code, Anthropic) for code generation, documentation "
    "drafting, and implementation acceleration. All technical decisions, experimental design choices, "
    "and reported results were supervised and verified by the human authors. No experimental data were "
    "fabricated; all tables and figures are derived from actual executed training runs. Statistical "
    "analysis and result interpretation reflect the authors' understanding of the methodology and "
    "findings."
)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved → {OUT}")
